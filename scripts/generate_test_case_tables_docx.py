"""Generate AutoGen_Test_Cases_Tables.docx for project report."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()


def main() -> None:
    out = Path(__file__).resolve().parent.parent / "AutoGen_Test_Cases_Tables.docx"

    doc = Document()
    t = doc.add_heading("AutoGen — Test case tables (report)", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Columns: Test case ID; Simple name (what it tests); Expected output; "
        "Actual output (per system design and backend code — replace with measured "
        "results after manual runs if required)."
    )

    headers = [
        "Test case ID",
        "Simple name (what it tests)",
        "Expected output",
        "Actual output (per system design / code)",
    ]

    # --- Authentication ---
    doc.add_heading("A. Authentication module", level=1)

    add_table(
        doc,
        "A1 — Black-box (10)",
        headers,
        [
            [
                "AUTH-BB-01",
                "Valid email/password sign-in",
                "HTTP 200; client receives tokens and user info",
                "POST /api/v1/auth/signin returns JSON with session.access_token, session.refresh_token, and user (id, email, user_metadata).",
            ],
            [
                "AUTH-BB-02",
                "Wrong email or password",
                "Clear failure; no session issued",
                "HTTP 401 with detail \"Invalid email or password.\" No session object in body.",
            ],
            [
                "AUTH-BB-03",
                "Sign-up with valid email and password (length rules)",
                "Account created or confirmation-required flag returned",
                "HTTP 200 with needs_email_confirmation true/false; if Supabase returns a session immediately, body also includes session and user.",
            ],
            [
                "AUTH-BB-04",
                "Sign-up with password shorter than 6 characters",
                "Request rejected before auth provider",
                "HTTP 422 (Pydantic): password min_length=6 validation error.",
            ],
            [
                "AUTH-BB-05",
                "Sign-up with invalid email format",
                "Request rejected before auth provider",
                "HTTP 422 (EmailStr validation failure on email).",
            ],
            [
                "AUTH-BB-06",
                "Token refresh with valid refresh token",
                "New access (and refresh) tokens returned",
                "HTTP 200 with new session and user from refresh_session.",
            ],
            [
                "AUTH-BB-07",
                "Token refresh with invalid or revoked refresh token",
                "Refresh denied",
                "HTTP 401 with detail \"Unable to refresh session.\"",
            ],
            [
                "AUTH-BB-08",
                "Get current user while signed in (valid access token)",
                "Current user and profile returned",
                "HTTP 200: {\"user\": ..., \"profile\": ...} after ensure_profile.",
            ],
            [
                "AUTH-BB-09",
                "Call protected API without Authorization header",
                "Access denied at gateway",
                "HTTP 403 from FastAPI HTTPBearer() (missing bearer).",
            ],
            [
                "AUTH-BB-10",
                "Get Google OAuth start URL (when configured)",
                "URL returned for browser redirect",
                "HTTP 200 with {\"url\": \"<oauth_url>\"}. If not configured, HTTP 400 with message that Google OAuth is not configured.",
            ],
        ],
    )

    add_table(
        doc,
        "A2 — White-box (10)",
        headers,
        [
            [
                "AUTH-WB-01",
                "get_current_user: Supabase AuthApiError on JWT",
                "Controlled 401, not generic 500",
                "Raises 401 with \"Session expired. Please sign in again.\"",
            ],
            [
                "AUTH-WB-02",
                "get_current_user: other exception during get_user",
                "Controlled 401",
                "Logs exception; raises 401 with \"Invalid or expired token.\"",
            ],
            [
                "AUTH-WB-03",
                "get_current_user: null user after get_user",
                "Controlled 401",
                "Raises 401 with \"Invalid or expired token.\"",
            ],
            [
                "AUTH-WB-04",
                "Successful sign_in after Supabase returns session",
                "Profile row ensured for user",
                "Calls ensure_profile(user.id, email, metadata) before returning tokens.",
            ],
            [
                "AUTH-WB-05",
                "Successful sign_up when session is immediately present",
                "Profile row ensured",
                "Same ensure_profile path as sign-in when result.session is not None.",
            ],
            [
                "AUTH-WB-06",
                "sign_up when Supabase returns user but no session",
                "No tokens in response",
                "Response contains needs_email_confirmation; session/user only when session exists.",
            ],
            [
                "AUTH-WB-07",
                "sign_up payload: full_name optional length cap",
                "Oversized name rejected at validation",
                "full_name optional, max_length=120 on AuthSignUpRequest; violation → 422.",
            ],
            [
                "AUTH-WB-08",
                "sign_in / sign_up password upper bound",
                "Oversized password rejected at validation",
                "Password max_length=128 on both auth request models; violation → 422.",
            ],
            [
                "AUTH-WB-09",
                "GET /api/v1/auth/google-url redirect target",
                "Redirect uses configured frontend base",
                "redirect_to set to settings.frontend_url + \"/login\" in OAuth options.",
            ],
            [
                "AUTH-WB-10",
                "POST /api/v1/auth/signout",
                "Server acknowledges sign-out",
                "Always returns {\"success\": true}; does not call Supabase server-side revoke in this handler (client clears session in UI).",
            ],
        ],
    )

    # --- OpenAI + PlantUML ---
    doc.add_heading("B. OpenAI + PlantUML call module", level=1)
    doc.add_paragraph(
        "End-to-end analysis pipeline: OpenAI for text/structure, PlantUML server for SVG."
    )

    add_table(
        doc,
        "B1 — Black-box (10)",
        headers,
        [
            [
                "AI-BB-01",
                "Run user-story analysis while logged in, server configured",
                "Project completes; ID returned",
                "HTTP 200 with {\"project_id\": \"<uuid>\", \"status\": \"completed\"} when pipeline succeeds.",
            ],
            [
                "AI-BB-02",
                "Run analysis when server has no OpenAI key",
                "User-visible \"unavailable\", no AI run",
                "HTTP 503 with detail that OPENAI_API_KEY is not configured (require_openai).",
            ],
            [
                "AI-BB-03",
                "Run user-story analysis with invalid JWT",
                "No analysis started",
                "401 from get_current_user before business logic.",
            ],
            [
                "AI-BB-04",
                "User-story payload: description too short",
                "Rejected by API validation",
                "422 from Pydantic: description min_length=20.",
            ],
            [
                "AI-BB-05",
                "User-story payload: project name too short",
                "Rejected by API validation",
                "422: project_name min_length=2.",
            ],
            [
                "AI-BB-06",
                "Code upload with zero usable files",
                "Clear client error",
                "400 \"At least one file is required.\" or \"Uploaded files were empty.\" depending on case.",
            ],
            [
                "AI-BB-07",
                "Successful code-upload analysis",
                "Project completes; tests later visible in result",
                "200 with project_id and completed; GET result exposes whitebox_tests / blackbox_tests.",
            ],
            [
                "AI-BB-08",
                "Mid-pipeline failure (OpenAI/PlantUML/storage)",
                "Project marked failed; user gets service error",
                "Project status failed; client receives 502 with message to check key, PlantUML, and Storage.",
            ],
            [
                "AI-BB-09",
                "Fetch another user's project_id on analyze result",
                "No data leak",
                "404 \"Project not found.\" (ownership filter on user_id).",
            ],
            [
                "AI-BB-10",
                "List projects when logged in",
                "Only own projects returned",
                "200 JSON array from projects filtered by user_id of current user.",
            ],
        ],
    )

    add_table(
        doc,
        "B2 — White-box (10)",
        headers,
        [
            [
                "AI-WB-01",
                "require_openai() when key missing",
                "Hard stop before Supabase project insert",
                "Raises 503 before user-story / code-upload inserts.",
            ],
            [
                "AI-WB-02",
                "User-story: primary AI bundle",
                "Diagram sources + summary from OpenAI",
                "Calls generate_user_story_bundle, then normalize_plantuml_block on architecture, use case, sequence.",
            ],
            [
                "AI-WB-03",
                "User-story: repair path",
                "Invalid PlantUML corrected when detector says so",
                "If is_plausible_plantuml false, calls repair_plantuml per diagram kind.",
            ],
            [
                "AI-WB-04",
                "User-story: PlantUML rendering pattern",
                "Parallel + single fetches",
                "fetch_two_svgs_parallel for architecture + use case; fetch_svg_bytes for sequence.",
            ],
            [
                "AI-WB-05",
                "User-story: persistence order",
                "DB reflects success only after storage",
                "Inserts analysis_results, uploads three SVGs under {userId}/{projectId}/diagrams/, then sets project completed.",
            ],
            [
                "AI-WB-06",
                "User-story: exception handling vs HTTPException",
                "Intentional HTTP errors not turned into failed wrongly",
                "HTTPException is re-raised; generic failures set failed and 502.",
            ],
            [
                "AI-WB-07",
                "Code-upload: diagram generation",
                "Control-flow + class from OpenAI",
                "generate_code_diagrams_bundle, normalize + optional repair_plantuml, fetch_two_svgs_parallel for both SVGs.",
            ],
            [
                "AI-WB-08",
                "Code-upload: test persistence",
                "One DB row per generated test",
                "Loops whitebox_tests and blackbox_tests with test_cases insert per item.",
            ],
            [
                "AI-WB-09",
                "Code-upload: source file storage attempt",
                "Files stored under deterministic prefix",
                "Upload path {userId}/{projectId}/{fileName} with upsert true; insert project_files metadata.",
            ],
            [
                "AI-WB-10",
                "Code-upload: swallow upload error in loop",
                "Upload failure does not always abort file row loop",
                "except pass around storage.upload — upload can fail silently for a file while loop continues.",
            ],
        ],
    )

    # --- Supabase storage ---
    doc.add_heading("C. Supabase storage module", level=1)
    doc.add_paragraph(
        "Private bucket: uploads, diagram SVGs, signed URLs, delete on project removal."
    )

    add_table(
        doc,
        "C1 — Black-box (10)",
        headers,
        [
            [
                "STG-BB-01",
                "Open completed analysis result in UI",
                "Diagram images can load",
                "API adds *_diagram_image_url fields via signed URL helper when paths exist.",
            ],
            [
                "STG-BB-02",
                "Completed user-story project",
                "Up to three diagram image URLs when stored",
                "URLs (or nulls) for architecture, use case, sequence path keys after enrichment.",
            ],
            [
                "STG-BB-03",
                "Completed code-upload project",
                "Two diagram image URLs when stored",
                "Control-flow and class diagram URL keys populated when paths exist.",
            ],
            [
                "STG-BB-04",
                "Delete a project from history",
                "Project and its artifacts logically gone",
                "200 {\"success\": true}; subsequent GETs for that id return 404.",
            ],
            [
                "STG-BB-05",
                "Unauthenticated client requests result",
                "No storage paths exposed",
                "403 missing bearer / 401 bad token — no result payload.",
            ],
            [
                "STG-BB-06",
                "Wrong user requests result for valid UUID",
                "No object listing for other tenant",
                "404 \"Project not found.\"",
            ],
            [
                "STG-BB-07",
                "Re-open result after some time",
                "New signed URLs when refetching",
                "Each GET .../result rebuilds signed URLs (1 hour in code) from current DB paths.",
            ],
            [
                "STG-BB-08",
                "Code upload with multiple files",
                "Analysis can still complete if storage OK",
                "Successful path stores sources and metadata; completed response when pipeline succeeds.",
            ],
            [
                "STG-BB-09",
                "Health endpoint",
                "Liveness without storage",
                "GET /health returns {\"status\": \"ok\"} with no storage call.",
            ],
            [
                "STG-BB-10",
                "Failed analysis project",
                "User may see failed status without new diagrams",
                "Project failed; typically no new completed diagram set for that run (depends on failure point).",
            ],
        ],
    )

    add_table(
        doc,
        "C2 — White-box (10)",
        headers,
        [
            [
                "STG-WB-01",
                "upload_svg upload options",
                "SVG stored as correct type, overwritable",
                "file_options content-type image/svg+xml and upsert true.",
            ],
            [
                "STG-WB-02",
                "create_signed_url when path is null",
                "No URL field crash",
                "Returns None immediately if storage_path is falsy.",
            ],
            [
                "STG-WB-03",
                "Signed URL TTL wired from API layer",
                "Predictable expiry",
                "enrich_analysis_result_with_signed_urls calls create_signed_url(..., 3600) (1 hour).",
            ],
            [
                "STG-WB-04",
                "create_signed_url on Supabase failure",
                "Graceful degradation",
                "Logs warning; returns None (frontend may show broken image for that slot).",
            ],
            [
                "STG-WB-05",
                "delete_paths with empty list",
                "No-op, no throw",
                "Function returns early when no non-empty paths.",
            ],
            [
                "STG-WB-06",
                "delete_paths strips blank paths",
                "Only real keys sent to remove",
                "Filters paths = [p for p in paths if p] before storage.remove.",
            ],
            [
                "STG-WB-07",
                "delete_paths on remove exception",
                "Delete does not crash API",
                "Logs warning \"storage delete failed\" and continues.",
            ],
            [
                "STG-WB-08",
                "Enrichment path-key mapping",
                "All diagram slots considered",
                "Loops five (path_key, url_key) pairs: architecture, use case, sequence, control flow, class.",
            ],
            [
                "STG-WB-09",
                "Code source storage.upload options",
                "Upsert on source objects",
                "upsert true for uploaded sources in create_code_upload_project.",
            ],
            [
                "STG-WB-10",
                "delete_project storage cleanup scope",
                "Deletes diagram + source objects",
                "Collects paths from analysis_results image columns and project_files.storage_path, then delete_paths before DB deletes.",
            ],
        ],
    )

    doc.save(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
